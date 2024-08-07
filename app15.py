

import streamlit as st
import pandas as pd
import google.generativeai as palm
from docx import Document
from io import BytesIO

# Load the Excel file
def a15():
    xls = pd.ExcelFile(r"Data.xlsx")

    # Load the sheet with the questions
    sheet_name ='15.Project Systems  '  # Adjust if necessary
    df = pd.read_excel(xls, sheet_name)

    # Example: Assume we identify correct column names for 'Category' and 'Qualification Question'

    # Assuming 'type' is the correct column name for 'Category'
    df['Category'] = df['Type '].fillna(method='ffill')

    # Drop rows where 'Program Discovery Questions' column is NaN
    questions_df = df.dropna(subset=['Question '])

    # Group questions by 'Category' and convert them to a dictionary
    questions_by_category = {
        category: questions['Question '].tolist()
        for category, questions in questions_df.groupby('Category')
    }


    # Initialize Google Generative AI
    palm.configure(api_key="AIzaSyBdb5z2W_YYIvyjk-iGN-DQY7uzcEyVEP4")

    # Streamlit app setup
    st.title("Project System Questionnaire")

    # Initialize session state variables
    if 'user_answers15' not in st.session_state:
        st.session_state.user_answers15 = {category: [""] * len(questions) for category, questions in questions_by_category.items()}

    if 'suggestions15' not in st.session_state:
        st.session_state.suggestions15 = {category: [""] * len(questions) for category, questions in questions_by_category.items()}

    def get_suggestion(question):
        try:
            # Use the Google Generative AI API to get a suggestion
            response = palm.generate_text(
                model='models/text-bison-001',  # Replace with the correct model name if necessary
                prompt=f"Provide a suggestion for the following question according to the project systems: {question}",
                max_output_tokens=150
            )
            suggestion = response.result
            return suggestion
        except Exception as e:
            return f"Error: {e}"

    # Sidebar for category selection
    st.sidebar.title("Categories")
    categories = list(questions_by_category.keys())
    selected_category = st.sidebar.radio("Select a category", options=categories + ["Summary"])

    # Display questions for the selected category
    if selected_category and selected_category != "Summary":
        st.session_state.selected_category = selected_category
        category = st.session_state.selected_category
        st.header(f"Category: {category}")
        questions = questions_by_category[category]
        
        for i, question in enumerate(questions):
            st.write(f"Q{i + 1}: {question}")
            
            user_answer = st.text_input("Your answer", value=st.session_state.user_answers15[category][i], key=f"answer_input_{category}_{i}")
            st.session_state.user_answers15[category][i] = user_answer
            
            if st.button("Suggestion", key=f"suggestion_{category}_{i}"):
                suggestion = get_suggestion(question)
                st.session_state.suggestions15[category][i] = suggestion
                st.experimental_rerun()
            
            # Display the suggestion if it exists
            if st.session_state.suggestions15[category][i]:
                st.write("Suggestion:")
                st.text_area("Suggested Answer", value=st.session_state.suggestions15[category][i], height=100, key=f"suggestion_text_{category}_{i}")

    if selected_category == "Summary":
        st.header("Summary of All Answers")
        summary_data = []
        
        for category, answers in st.session_state.user_answers15.items():
            for i, answer in enumerate(answers):
                question = questions_by_category[category][i]
                
                summary_data.append({'Category': category, 'Question': question, 'Answer': answer})
                st.write(f"Q{i + 1}: {question}")
                st.write(f"A{i + 1}: {answer}")
        summary_df = pd.DataFrame(summary_data)
        #st.write(summary_df)

        # Save to Excel button
    
    

        # Submit button for all answers with confirmation
            # Save to Excel button
        user_name = st.text_input("Enter your name:")
        if user_name:
            st.write(f"Hello, {user_name}!")
    

        ## Submit button for all answers with confirmatio
        if st.button("Submit"):
            try:
                if summary_df.shape[0] > 0:
                    # Save answers to a BytesIO buffer
                    excel_buffer = BytesIO()
                    summary_df.to_excel(excel_buffer, index=False)
                    excel_buffer.seek(0)

                    # Generate scope suggestion
                    summary_text = "\n".join([f"{row['Question']}: {row['Answer']}" for _, row in summary_df.iterrows()])
                    response = palm.generate_text(
                        model='models/text-bison-001',  # Adjust model name if necessary
                        prompt=f"Based on the following answers, what is the scope for SAP? generate paragraph in concise manner\n\n{summary_text}",
                        max_output_tokens=300
                    )
                    scope_suggestion = response.result

                    # Save the scope suggestion to a BytesIO buffer as a Word document
                    doc = Document()
                    doc.add_heading("Scope for SAP", level=1)
                    doc.add_paragraph(scope_suggestion)
                    word_buffer = BytesIO()
                    doc.save(word_buffer)
                    word_buffer.seek(0)

                    # Provide download buttons
                    st.download_button(
                        label="Download Answers",
                        data=excel_buffer,
                        file_name=f"{user_name}_answers.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

                    st.download_button(
                        label="Download Scope for SAP",
                        data=word_buffer,
                        file_name=f"{user_name}_scope_for_SAP.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.success("Files are ready for download.")
                else:
                    st.warning("No answers provided. Nothing to save.")
            except Exception as e:
                st.error(f"Error during submission: {e}")

        
                    


    # Submit button for all answers




