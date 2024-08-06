import streamlit as st
import pandas as pd
import google.generativeai as palm
import os
from docx import Document
from io import BytesIO

# Load the Excel file
def a3():
    xls = pd.ExcelFile(r"Data.xlsx")

    # Load the sheet with the questions
    sheet_name ='3.Landscape Details'  # Adjust if necessary
    df = pd.read_excel(xls, sheet_name)

    # Example: Assume we identify correct column names for 'Category' and 'Qualification Question'

    # Assuming 'type' is the correct column name for 'Category'
    df['Category'] = df['Component'].fillna(method='ffill')

    # Drop rows where 'Program Discovery Questions' column is NaN
    questions_df = df.dropna(subset=['Query'])

    # Group questions by 'Category' and convert them to a dictionary
    questions_by_category = {
        category: questions['Query'].tolist()
        for category, questions in questions_df.groupby('Category')
    }

    # Initialize Google Generative AI
    api_key = os.getenv("GOOGLE_API_KEY")

    # Streamlit app setup
    st.title("Landscape Questionnaire")

    # Initialize session state variables
    if 'user_answers1' not in st.session_state:
        st.session_state.user_answers1 = {category: [""] * len(questions) for category, questions in questions_by_category.items()}

    if 'suggestions1' not in st.session_state:
        st.session_state.suggestions1 = {category: [""] * len(questions) for category, questions in questions_by_category.items()}

    def get_suggestion(question):
        try:
            # Use the Google Generative AI API to get a suggestion
            response = palm.generate_text(
                model='models/text-bison-001',  # Replace with the correct model name if necessary
                prompt=f"Provide a suggestion for the following question according to the landscape of sap: {question}",
                max_output_tokens=150
            )
            suggestion = response.result
            return suggestion
        except Exception as e:
            return f"Error: {e}"

    # Sidebar for category selection
    st.sidebar.title("Categories")
    categories = list(questions_by_category.keys())
    selected_category = st.sidebar.radio("Select a category", options=categories + ["Summary"])

    # Display questions for the selected category
    if selected_category and selected_category != "Summary":
        st.session_state.selected_category = selected_category
        category = st.session_state.selected_category
        st.header(f"Category: {category}")
        questions = questions_by_category[category]
        
        for i, question in enumerate(questions):
            st.write(f"Q{i + 1}: {question}")
            
            user_answer = st.text_input("Your answer", value=st.session_state.user_answers1[category][i], key=f"answer_input_{category}_{i}")
            st.session_state.user_answers1[category][i] = user_answer
            
            if st.button("Suggestion", key=f"suggestion_{category}_{i}"):
                suggestion = get_suggestion(question)
                st.session_state.suggestions1[category][i] = suggestion
                st.experimental_rerun()
            
            # Display the suggestion if it exists
            if st.session_state.suggestions1[category][i]:
                st.write("Suggestion:")
                st.text_area("Suggested Answer", value=st.session_state.suggestions1[category][i], height=100, key=f"suggestion_text_{category}_{i}")

    # Display summary of all answers
    if selected_category == "Summary":
        st.header("Summary of All Answers")
        summary_data = []
        
        for category, answers in st.session_state.user_answers1.items():
            for i, answer in enumerate(answers):
                question = questions_by_category[category][i]
                
                summary_data.append({'Category': category, 'Question': question, 'Answer': answer})
                st.write(f"Q{i + 1}: {question}")
                st.write(f"A{i + 1}: {answer}")
        summary_df = pd.DataFrame(summary_data)

        # Convert summary_df to Excel file for download
        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
            summary_df.to_excel(writer, index=False, sheet_name='Summary')
        excel_buffer.seek(0)

        # Convert scope suggestion to Word document for download
        def create_word_document(text):
            doc = Document()
            doc.add_heading("Scope for SAP", level=1)
            doc.add_paragraph(text)
            word_buffer = BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)
            return word_buffer

        # User name input
        user_name = st.text_input("Enter your name:")
        if user_name:
            st.write(f"Hello, {user_name}!")

        if st.button("Submit"):
            try:
                # Check if there are answers to save
                if summary_df.shape[0] > 0:
                    # Prepare summary text
                    summary_text = "\n".join([f"{row['Question']}: {row['Answer']}" for _, row in summary_df.iterrows()])
                    # Ask the chatbot "What is the scope for SAP"
                    response = palm.generate_text(
                        model='models/text-bison-001',  # Adjust model name if necessary
                        prompt=f"Based on the following answers, what is the scope for SAP? Generate paragraph in concise manner\n\n{summary_text}",
                        max_output_tokens=300
                    )
                    scope_suggestion = response.result

                    # Create word document buffer
                    word_buffer = create_word_document(scope_suggestion)

                    # Provide download buttons
                    st.download_button(
                        label="Download Answers",
                        data=excel_buffer,
                        file_name=f"{user_name}_answers.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

                    st.download_button(
                        label="Download Scope for SAP",
                        data=word_buffer,
                        file_name=f"{user_name}_scope_for_SAP.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success(f"Answers and Scope for SAP are ready for download.")
                else:
                    st.warning("No answers provided. Nothing to save.")
            except Exception as e:
                st.error(f"Error during submission: {e}")

# Run the app
a3()





