import streamlit as st
import pandas as pd
import google.generativeai as palm
import os
from docx import Document
from io import BytesIO

def a1():
    # Load the Excel file
    file_path = "Data.xlsx"
    xls = pd.ExcelFile(file_path)

    # Load the sheet with the questions
    sheet_name = xls.sheet_names[0]  # Adjust if necessary
    df = pd.read_excel(xls, sheet_name)

    # Extract categories and questions from the DataFrame (adjust column names as needed)
    df['Category'] = df['Category'].fillna(method='ffill')
    categories = df['Category'].unique()

    # Create a dictionary of questions by category
    questions_by_category = {category: df[df['Category'] == category]['Qualification Question'].dropna().tolist() for category in categories}

    # Initialize Google Generative AI
    palm.configure(api_key="AIzaSyBdb5z2W_YYIvyjk-iGN-DQY7uzcEyVEP4")
    
    # Streamlit app setup
    st.title("Qualification Questionnaire")

    # Initialize session state variables
    if 'user_answers' not in st.session_state:
        st.session_state.user_answers = {category: [""] * len(questions) for category, questions in questions_by_category.items()}

    if 'suggestions' not in st.session_state:
        st.session_state.suggestions = {category: [""] * len(questions) for category, questions in questions_by_category.items()}

    def get_suggestion(question):
        try:
            # Use the Google Generative AI API to get a suggestion
            response = palm.generate_text(
                model='models/text-bison-001',  # Adjust model name if necessary
                prompt=f"Provide a suggestion for the following question: {question}",
                max_output_tokens=150
            )
            suggestion = response.result
            return suggestion
        except Exception as e:
            st.error(f"Error: {e}")
            return None

    # Sidebar for category selection
    st.sidebar.title("Categories")
    categories = list(questions_by_category.keys())
    selected_category = st.sidebar.radio("Select a category", options=categories + ["Summary"])

    # Display questions for the selected category
    if selected_category and selected_category != "Summary":
        st.session_state.selected_category = selected_category
        category = st.session_state.selected_category
        st.header(f"Category: {category}")
        questions = questions_by_category[category]
        
        for i, question in enumerate(questions):
            st.write(f"Q{i + 1}: {question}")

            # Check if the question is Prospect Name & Address
            if "Prospect Name & Address" in question:
                # Split the question into Prospect Name and Address
                sub_questions = ["Prospect Name", "Address"]
                if not isinstance(st.session_state.user_answers[category][i], dict):
                    st.session_state.user_answers[category][i] = {'Prospect Name': "", 'Address': ""}
                
                for sub_question in sub_questions:
                    user_answer = st.text_input(sub_question, value=st.session_state.user_answers[category][i][sub_question], key=f"answer_input_{category}_{i}_{sub_question}")
                    st.session_state.user_answers[category][i][sub_question] = user_answer
            else:
                # Handling specific question
                if "Please indicate the business areas which you would like to address in your evaluation" in question:
                    options = ["Marketing", "Sales", "Project management", "Finance", "Accounting", "Human Resources", "Purchasing", "Inventory", "Manufacturing", "Quality"]
                    selected_options = []
                    for option in options:
                        selected = st.checkbox(option, key=f"checkbox_{category}_{i}_{option}")
                        if selected:
                            selected_options.append(option)
                    st.session_state.user_answers[category][i] = selected_options
                    
                else:
                    # Text input for user's answer
                    user_answer = st.text_input("Your answer", value=st.session_state.user_answers[category][i], key=f"answer_input_{category}_{i}")
                    st.session_state.user_answers[category][i] = user_answer

                    # Suggestion button
                    if not any(exclude_word in question for exclude_word in ["Prospect Name:", "Address:", "Please indicate the business areas which you would like to address in your evaluation:"]):
                        if st.button("Suggestion", key=f"suggestion_{category}_{i}"):
                            suggestion = get_suggestion(question)
                            if suggestion:
                                st.session_state.suggestions[category][i] = suggestion
                            else:
                                st.warning("Failed to generate suggestion. Please try again.")

                        # Display suggestion in a separate box
                        if st.session_state.suggestions[category][i]:
                            st.text_area("Suggestion", value=st.session_state.suggestions[category][i], height=100, key=f"suggestion_box_{category}_{i}")

    # Display summary of all answers
    if selected_category == "Summary":
        st.header("Summary of All Answers")
        summary_data = []
        
        for category, answers in st.session_state.user_answers.items():
            for i, answer in enumerate(answers):
                question = questions_by_category[category][i]
                if "Prospect Name & Address" in question:
                    if not isinstance(st.session_state.user_answers[category][i], dict):
                        st.session_state.user_answers[category][i] = {'Prospect Name': "", 'Address': ""}
                    prospect_name_answer = st.session_state.user_answers[category][i]['Prospect Name']
                    address_answer = st.session_state.user_answers[category][i]['Address']
                    summary_data.append({'Category': category, 'Question': 'Prospect Name', 'Answer': prospect_name_answer})
                    summary_data.append({'Category': category, 'Question': 'Address', 'Answer': address_answer})
                    st.write(f"Q{i + 1}: Prospect Name")
                    st.write(f"A{i + 1}: {prospect_name_answer}")
                    
                    st.write(f"Q{i + 2}: Address")
                    st.write(f"A{i + 2}: {address_answer}")
                else:
                    summary_data.append({'Category': category, 'Question': question, 'Answer': answer})
                    st.write(f"Q{i + 1}: {question}")
                    st.write(f"A{i + 1}: {answer}")
        summary_df = pd.DataFrame(summary_data)

        # Save to Excel button
        user_name = st.text_input("Enter your name:")
        if user_name:
            st.write(f"Hello, {user_name}!")
    
        # Submit button for all answers with confirmation
        if st.button("Submit"):
            try:
                # Check if there are answers to save
                if summary_df.shape[0] > 0:
                    # Save answers to Excel file
                    excel_buffer = BytesIO()
                    summary_df.to_excel(excel_buffer, index=False)
                    excel_buffer.seek(0)
                    
                    # Prepare summary text
                    summary_text = "\n".join([f"{row['Question']}: {row['Answer']}" for _, row in summary_df.iterrows()])
                    # Ask the chatbot "What is the scope for SAP"
                    response = palm.generate_text(
                        model='models/text-bison-001',  # Adjust model name if necessary
                        prompt=f"Based on the following answers, what is the scope for SAP?generate paragraph in concise manner\n\n{summary_text}",
                        max_output_tokens=300
                    )
                    scope_suggestion = response.result

                    # Save the scope suggestion to a Word document
                    doc = Document()
                    doc.add_heading("Scope for SAP", level=1)
                    doc.add_paragraph(scope_suggestion)
                    word_buffer = BytesIO()
                    doc.save(word_buffer)
                    word_buffer.seek(0)

                    # Provide download links for the generated files
                    st.download_button(
                        label="Download Answers",
                        data=excel_buffer,
                        file_name=f"{user_name}_answers.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

                    st.download_button(
                        label="Download Scope for SAP",
                        data=word_buffer,
                        file_name=f"{user_name}_scope_for_SAP.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    st.success(f"Answers and Scope for SAP are ready for download.")
                else:
                    st.warning("No answers provided. Nothing to save.")
            except Exception as e:
                st.error(f"Error during submission: {e}")

if __name__ == "__main__":
    a1()




