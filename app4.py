import streamlit as st
import pandas as pd
import google.generativeai as palm
from docx import Document
import io

def a4():
    # Load the Excel file
    file_path = r"Data.xlsx"
    xls = pd.ExcelFile(file_path)

    # List all sheet names
    sheet_names = xls.sheet_names

    # Given sheet name
    sheet_name = '4.Organisation Structure'

    # Function to check if sheet exists
    def check_sheet_exists(sheet_name, sheet_names):
        if sheet_name not in sheet_names:
            raise ValueError(f"The sheet '{sheet_name}' is not present in the Excel file. Available sheets are: {', '.join(sheet_names)}")

    # Check if the given sheet exists and handle the exception
    try:
        check_sheet_exists(sheet_name, sheet_names)
        
        df = pd.read_excel(xls, sheet_name)
    except ValueError as e:
        st.error(str(e))
        st.stop()

    # Drop rows where 'Question' column is NaN
    questions_df = df.dropna(subset=['Questions'])

    # Convert questions to a list
    questions = questions_df['Questions'].tolist()

    # Initialize Google Generative AI
    palm.configure(api_key="AIzaSyBdb5z2W_YYIvyjk-iGN-DQY7uzcEyVEP4") # Replace with your API key

    # Streamlit app setup
    st.title("Organisation Structure Questionnaire")

    # Initialize session state variables
    if 'user_answers3' not in st.session_state:
        st.session_state.user_answers3 = [""] * len(questions)

    if 'suggestions3' not in st.session_state:
        st.session_state.suggestions3 = [""] * len(questions)

    def get_suggestion(question):
        try:
            # Use the Google Generative AI API to get a suggestion
            response = palm.generate_text(
                model='models/text-bison-001',  # Replace with the correct model name if necessary
                prompt=f"Provide a suggestion for the following question according to the organization structure: {question}",
                max_output_tokens=150
            )
            suggestion = response.result
            return suggestion
        except Exception as e:
            return f"Error: {e}"

    # Sidebar for navigation
    st.sidebar.title("Navigation")
    selected_option = st.sidebar.radio("Go to", options=["Questionnaire", "Summary"])

    if selected_option == "Questionnaire":
        # Display questions
        for i, question in enumerate(questions):
            st.write(f"Q{i + 1}: {question}")

            user_answer = st.text_input("Your answer", value=st.session_state.user_answers3[i], key=f"answer_input_{i}")
            st.session_state.user_answers3[i] = user_answer

            # Button to generate suggestion
            if st.button("Suggestion", key=f"suggestion_{i}"):
                suggestion = get_suggestion(question)
                st.session_state.suggestions3[i] = suggestion
                st.experimental_rerun()

            # Display the suggestion if it exists
            if st.session_state.suggestions3[i]:
                st.write("Suggestion:")
                st.text_area("Suggested Answer", value=st.session_state.suggestions3[i], height=100, key=f"suggestion_text_{i}")

    elif selected_option == "Summary":
        # Display summary of all answers
        st.header("Summary of All Answers")
        summary_data = []

        for i, question in enumerate(questions):
            answer = st.session_state.user_answers3[i]
            summary_data.append({'Question': question, 'Answer': answer})
            st.write(f"Q{i + 1}: {question}")
            st.write(f"A{i + 1}: {answer}")

        summary_df = pd.DataFrame(summary_data)

        # Generate Excel file in memory
        excel_io = io.BytesIO()
        summary_df.to_excel(excel_io, index=False)
        excel_io.seek(0)
        excel_data = excel_io.getvalue()

        # Generate scope suggestion
        user_name = st.text_input("Enter your name:")
        if user_name:
            st.write(f"Hello, {user_name}!")

        # Submit button for all answers with confirmation
        if st.button("Submit"):
            try:
                # Check if there are answers to save
                if summary_df.shape[0] > 0:
                    # Prepare summary text
                    summary_text = "\n".join([f"{row['Question']}: {row['Answer']}" for _, row in summary_df.iterrows()])
                    # Ask the chatbot "What is the scope for SAP"
                    response = palm.generate_text(
                        model='models/text-bison-001',  # Adjust model name if necessary
                        prompt=f"Based on the following answers, what is the scope for SAP? generate paragraph in concise manner\n\n{summary_text}",
                        max_output_tokens=300
                    )
                    scope_suggestion = response.result
                    
                    # Save the scope suggestion to a Word document in memory
                    doc = Document()
                    doc.add_heading("Scope for SAP", level=1)
                    doc.add_paragraph(scope_suggestion)
                    word_io = io.BytesIO()
                    doc.save(word_io)
                    word_io.seek(0)
                    word_data = word_io.getvalue()

                    # Create download links
                    st.download_button(
                        label="Download Your Answers",
                        data=excel_data,
                        file_name=f"{user_name}_answers.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

                    st.download_button(
                        label="Download Scope for SAP",
                        data=word_data,
                        file_name=f"{user_name}_scope_for_SAP.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    st.success(f"Answers and Scope for SAP are ready for download.")
                else:
                    st.warning("No answers provided. Nothing to save.")
            except Exception as e:
                st.error(f"Error during submission: {e}")

# Run the app
if __name__ == "__main__":
    a4()

